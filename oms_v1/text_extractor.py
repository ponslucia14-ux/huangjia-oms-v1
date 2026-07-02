from __future__ import annotations

import shutil
import subprocess
from pathlib import Path

from docx import Document
from pypdf import PdfReader

from .schemas import InputEnvelope, TextExtraction


IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp"}


class TextExtractor:
    """Extract text from direct text, documents, PDFs, or OCR-capable images."""

    def extract(self, envelope: InputEnvelope) -> TextExtraction:
        if envelope.text.strip():
            return TextExtraction(status="ok", text=envelope.text, source="envelope_text")

        if not envelope.file_path:
            return TextExtraction(status="empty", text="", source="none", error="no text or file_path")

        path = Path(envelope.file_path)
        if not path.exists():
            return TextExtraction(status="error", text="", source="file", error=f"file not found: {path}")

        suffix = path.suffix.lower()
        try:
            if suffix in {".txt", ".md", ".csv"}:
                return TextExtraction(status="ok", text=path.read_text(encoding="utf-8-sig"), source="file_text")
            if suffix == ".docx":
                return TextExtraction(status="ok", text=self._extract_docx(path), source="docx")
            if suffix == ".pdf":
                return TextExtraction(status="ok", text=self._extract_pdf(path), source="pdf")
            if suffix in IMAGE_SUFFIXES:
                return self._extract_image(path)
        except Exception as exc:
            return TextExtraction(status="error", text="", source=suffix.lstrip("."), error=repr(exc))

        return TextExtraction(status="unsupported", text="", source="file", error=f"unsupported file type: {suffix}")

    def _extract_docx(self, path: Path) -> str:
        doc = Document(path)
        parts: list[str] = []
        parts.extend(p.text.strip() for p in doc.paragraphs if p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " / ") for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    def _extract_pdf(self, path: Path) -> str:
        reader = PdfReader(str(path))
        return "\n".join((page.extract_text() or "").strip() for page in reader.pages).strip()

    def _extract_image(self, path: Path) -> TextExtraction:
        sidecar = path.with_suffix(path.suffix + ".txt")
        if sidecar.exists():
            return TextExtraction(status="ok", text=sidecar.read_text(encoding="utf-8-sig"), source="image_sidecar_text")

        tesseract = shutil.which("tesseract")
        if not tesseract:
            return TextExtraction(
                status="ocr_unavailable",
                text="",
                source="image",
                engine="tesseract_cli",
                error="tesseract command not found; provide sidecar .txt or install OCR engine",
            )

        cmd = [tesseract, str(path), "stdout", "-l", "chi_sim+eng"]
        proc = subprocess.run(cmd, capture_output=True, text=True, timeout=60, check=False)
        if proc.returncode != 0:
            return TextExtraction(
                status="error",
                text="",
                source="image",
                engine="tesseract_cli",
                error=proc.stderr.strip() or f"tesseract failed: {proc.returncode}",
            )
        return TextExtraction(status="ok", text=proc.stdout.strip(), source="image_ocr", engine="tesseract_cli")

